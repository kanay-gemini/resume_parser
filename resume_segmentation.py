import docx
import docx2txt
import os
import re
import argparse


# parser = argparse.ArgumentParser()
# parser.add_argument('filename', type=str, help='give the complete path of the resume')
# args = parser.parse_args()
# filename = os.path.join(os.getcwd(), '{}'.format(args.filename))

# def document_object(filename):
#     document = docx.Document(filename)
#     return document

# document = document_object(filename)


def document_object(filename):
    document = docx.Document(filename)
    return document


def properties_identifier(document):
    '''
    This method is responsible for extracting various properties from document object 
    and classify them as headings, bold, italic, font color, font size etc
    '''
    bold = []
    h1 = []
    h2 = []
    h3 = []
    h4 = []
    italic = []
    color_dict = {}
    color_size_dict={}
    font_dict = {}
    font_size_dict = {}
    try:
        level_from_style_name = {f'Heading {i}': f'<H{i}>' for i in range(10)}

        for para in document.paragraphs:
            if para.style.name in level_from_style_name:
                level = level_from_style_name[para.style.name]
                if level == "<H1>":
                    text=para.text
                    new=""
                    for char in text:
                        if char.isalpha() or char==" ":
                            new+=char
                    h1.append(new)
                if level == "<H2>":
                    h2.append(para.text)
                if level == "<H3>":
                    h3.append(para.text)
                if level == "<H4>":
                    h4.append(para.text)

            for run in para.runs:
                if run.bold:
                    text=run.text
                    new=""
                    for char in text:
                        if char.isalpha() or char==" ":
                            new+=char
                    bold.append(new)
                if run.italic:
                    italic.append(run.text)

            for run in para.runs:
                if run.font.color.rgb is not None:
                    if run.font.color.rgb in color_dict:
                        color_dict[run.font.color.rgb].append(run.text)
                    else:
                        color_dict[run.font.color.rgb] = [run.text]

                if run.font.color.rgb is not None:
                    if run.font.color.rgb in color_size_dict:
                        color_size_dict[run.font.color.rgb]+=1
                    else:
                        color_size_dict[run.font.color.rgb]=1

                if run.font.size is not None:
                    if run.font.size in font_dict:
                        font_dict[run.font.size].append(run.text)
                    else:
                        font_dict[run.font.size] = [run.text]

                if run.font.size is not None:
                    if run.font.size in font_size_dict:
                        font_size_dict[run.font.size] += 1
                    else:
                        font_size_dict[run.font.size] = 1
        result = (bold, italic, h1, h2, h3, h4, color_dict, color_size_dict, font_dict, font_size_dict)
        return result
    except Exception as e:
        print("error occured in extracting properties from document object")
        result = (bold, italic, h1, h2, h3, h4, color_dict, color_size_dict, font_dict, font_size_dict)
        return result


# HEADERS_OF_RESUME = []
def bold_headings(bold, synonym_dict):
    no_of_headings_in_bold = 0
    try:
        for heading in synonym_dict["main_headings"]:
            if heading in [i.lower() for i in bold]:
                no_of_headings_in_bold += 1
        return no_of_headings_in_bold
    except Exception as e:
        return no_of_headings_in_bold


def h1_headings(h1, synonym_dict):
    try:
        no_of_headings_in_h1 = 0
        for heading in synonym_dict["main_headings"]:
            if heading in [i.lower() for i in h1]:
                no_of_headings_in_h1 += 1
        return no_of_headings_in_h1
    except Exception as e:
        return no_of_headings_in_h1


def h2_headings(h2, synonym_dict):
    try:
        no_of_headings_in_h2 = 0
        for heading in synonym_dict["main_headings"]:
            if heading in [i.lower() for i in h2]:
                no_of_headings_in_h2 += 1
        return no_of_headings_in_h2
    except Exception as e:
        return no_of_headings_in_h2


def h3_headings(h3, synonym_dict):
    try:
        no_of_headings_in_h3 = 0
        for heading in synonym_dict["main_headings"]:
            if heading in [i.lower() for i in h3]:
                no_of_headings_in_h3 += 1
        return no_of_headings_in_h3
    except Exception as e:
        return no_of_headings_in_h3


count_size = 0
count_color = 0
font_dict_header = []
color_dict_header = []


def font_size(font_dict, synonym_dict):
    no_of_headings_in_font_size = 0
    font_dict_header = []
    count_size = 0
    try:
        for values_list in list(font_dict.values()):
            for heading in synonym_dict["main_headings"]:
                if heading in [i.lower() for i in values_list]:
                    no_of_headings_in_font_size += 1
                    if no_of_headings_in_font_size > count_size:
                        count_size = no_of_headings_in_font_size
                        font_dict_header = values_list
        return (no_of_headings_in_font_size, font_dict_header)
    except Exception as e:
        return (no_of_headings_in_font_size, font_dict_header)


def font_color(color_dict, synonym_dict):
    no_of_headings_in_color = 0
    color_dict_header = []
    count_color = 0
    try:
        for values_list in list(color_dict.values()):
            for heading in synonym_dict["main_headings"]:
                if heading in [i.lower() for i in values_list]:
                    no_of_headings_in_color += 1
                    if no_of_headings_in_color > count_color:
                        count_color = no_of_headings_in_color
                        color_dict_header = values_list
        return (no_of_headings_in_color, color_dict_header)
    except Exception as e:
        return (no_of_headings_in_color, color_dict_header)


def maximum_headings(bold_headings , h1_headings, h2_headings, h3_headings, count_size, count_color):
    header_list_count = [bold_headings, h1_headings,
                     h2_headings, h3_headings, count_size,count_color]
    maximum = max(header_list_count)
    return maximum

def resume_headings(maximum, bold_headings, h1_headings, h2_headings, h3_headings, count_size, count_color, bold, h1, h2, h3, font_dict_header, color_dict_header):
    HEADERS_OF_RESUME = []
    if bold_headings == maximum:
        HEADERS_OF_RESUME = bold
    elif h1_headings == maximum:
        HEADERS_OF_RESUME = h1
    elif h2_headings == maximum:
        HEADERS_OF_RESUME = h2
    elif h3_headings == maximum:
        HEADERS_OF_RESUME = h3
    elif count_size == maximum:
        HEADERS_OF_RESUME = font_dict_header
    elif count_color == maximum:
        HEADERS_OF_RESUME = color_dict_header
    else:
        pass
    return HEADERS_OF_RESUME


def extract_text_from_doc(filename):
    temp = docx2txt.process(filename)
    text = [line.replace('\t', ' ') for line in temp.split('\n') if line]
    return '\n'.join(text)


def customized_headers(HEADERS_OF_RESUME, synonym_dict, complete_page_text):
    NEW_HEADERS = []
    headers=0

    for heading in HEADERS_OF_RESUME:
        heading = heading.strip()
        if len(heading) > 3:
            if heading != '' and heading != '\t' and heading != ':':
                NEW_HEADERS.append(heading)
    HEADERS_OF_RESUME = NEW_HEADERS

    try:
        for header in HEADERS_OF_RESUME:
            if header.lower() in synonym_dict["main_headings"]:
                headers+=1
        
        if headers<3:
            all_lines=complete_page_text.split('\n')
            NEW_HEADERS=[]
            for header in all_lines:
                if all(chr.isalpha() or chr.isspace() for chr in header) and len(header.split())<=3 and len(header.split())>=1:
                    header=header.strip()
                    NEW_HEADERS.append(header)
            HEADERS_OF_RESUME=NEW_HEADERS
            LAST=[]
            for header in HEADERS_OF_RESUME:
                if header.lower() in synonym_dict["main_headings"]:
                    LAST.append(header)
            HEADERS_OF_RESUME=LAST
            return HEADERS_OF_RESUME
        
        return HEADERS_OF_RESUME
    except Exception as e:
        print("error occured in cutomizing headers")
        HEADERS_OF_RESUME = NEW_HEADERS
        return HEADERS_OF_RESUME


def text_between_various_headings(complete_page_text, synonym_dict, HEADERS_OF_RESUME):
    career_objective_segment = ''
    personal_segment = ''
    education_segment = ''
    professional_segment = ''
    skills_segment = ''
    try:
        length = len(HEADERS_OF_RESUME)
        flag=0
        for header in HEADERS_OF_RESUME:
            if header.lower() in synonym_dict["main_headings"]:
                personal_segment = complete_page_text.split(header)[0]
                break
        if personal_segment:
            flag=1

        for i in range(length-1):
            if HEADERS_OF_RESUME[i].lower() in synonym_dict["personal details"]:
                if flag==1:
                    personal_segment+=complete_page_text.split(
                    HEADERS_OF_RESUME[i])[1].split(HEADERS_OF_RESUME[i+1])[0]
                else:
                    personal_segment = complete_page_text.split(
                    HEADERS_OF_RESUME[i])[1].split(HEADERS_OF_RESUME[i+1])[0]
            elif HEADERS_OF_RESUME[i].lower() in synonym_dict["career_objective"]:
                career_objective_segment = complete_page_text.split(
                    HEADERS_OF_RESUME[i])[1].split(HEADERS_OF_RESUME[i+1])[0]
            elif HEADERS_OF_RESUME[i].lower() in synonym_dict["education"]:
                education_segment = complete_page_text.split(
                    HEADERS_OF_RESUME[i])[1].split(HEADERS_OF_RESUME[i+1])[0]
            elif HEADERS_OF_RESUME[i].lower() in synonym_dict["professional_experience"]:
                professional_segment = complete_page_text.split(
                    HEADERS_OF_RESUME[i])[1].split(HEADERS_OF_RESUME[i+1])[0]
            elif HEADERS_OF_RESUME[i].lower() in synonym_dict["skills"]:
                skills_segment = complete_page_text.split(
                    HEADERS_OF_RESUME[i])[1].split(HEADERS_OF_RESUME[i+1])[0]
            else:
                pass
        
        if((i==length-2) and (HEADERS_OF_RESUME[length-1].lower() in synonym_dict["career_objective"])):
            career_objective_segment =complete_page_text.split(HEADERS_OF_RESUME[length-1])[1]
        elif((i==length-2) and (HEADERS_OF_RESUME[length-1].lower() in synonym_dict["education"])):
            education_segment =complete_page_text.split(HEADERS_OF_RESUME[length-1])[1]
        elif((i==length-2) and (HEADERS_OF_RESUME[length-1].lower() in synonym_dict["professional_experience"])):
            professional_segment =complete_page_text.split(HEADERS_OF_RESUME[length-1])[1]
        elif((i==length-2) and (HEADERS_OF_RESUME[length-1].lower() in synonym_dict["skills"])):
            skills_segment =complete_page_text.split(HEADERS_OF_RESUME[length-1])[1]
        elif((i==length-2) and (HEADERS_OF_RESUME[length-1].lower() in synonym_dict["personal details"])):
            if flag==1:
                personal_segment+=complete_page_text.split(
                HEADERS_OF_RESUME[length-1])[1]
            else:
                personal_segment = complete_page_text.split(
                HEADERS_OF_RESUME[length-1])[1]
        return (career_objective_segment, personal_segment, education_segment, skills_segment, professional_segment)

    except Exception as e:
        print("error occured in extracting text between various headings")
        return (career_objective_segment, personal_segment, education_segment, skills_segment, professional_segment)
